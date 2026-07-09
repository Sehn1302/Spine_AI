"""Build LinkedIn post Word document (local only — not for GitHub)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "marketing" / "Spine_AI_LinkedIn_Post.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("Spine AI — LinkedIn Post", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Ready to copy into LinkedIn. Hashtags and X version at the end.")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True

    doc.add_paragraph()

    add_heading(doc, "Main post", 1)

    paragraphs = [
        "ChatGPT, Copilot, and enterprise assistants all share the same hidden trade-off: "
        "your data goes to their cloud, their model does every job, and your PC is mostly an afterthought.",
        "I did not want another tab in a browser. I wanted an assistant that lives on my machine — "
        "listens, thinks, remembers, and acts — without sending my voice, files, or commands to a third party.",
        "So I built Spine AI.",
        "Not as a chatbot clone. As a local AI operating layer: multiple models, multiple specialist agents, "
        "one orchestrator brain — voice in, voice out, with real control over the computer it runs on.",
    ]
    for p in paragraphs:
        doc.add_paragraph(p)

    add_heading(doc, "How I built it — vibe coding + AI product engineering", 2)
    doc.add_paragraph(
        "I used AI-assisted development (“vibe coding”) to move fast — describing outcomes in natural language "
        "and iterating with Cursor — but I owned the product and system design: multi-model routing, agent "
        "architecture, voice pipeline, GPU optimisation, stress testing, and UX (orb, tray, boot startup). "
        "That is AI product engineering: directing AI to build software while designing how the whole system behaves."
    )

    add_heading(doc, "What I did", 2)
    add_bullets(
        doc,
        [
            "Designed a multi-LLM orchestrator that routes tasks to the right model — not one model for everything",
            "Built specialist agents (research, study, files, PC control) under a single “brain”",
            "Shipped voice-to-voice on local hardware: GPU Whisper (CUDA 12) + Edge TTS",
            "Added RAG memory — Spine recalls your indexed notes and files",
            "Enabled natural-language PC control — “write this down”, organise files, open apps, Spotify, research papers",
            "Stress-tested routing, concurrency, and model switching (11/11 checks passed)",
        ],
    )

    add_heading(doc, "How Spine is different from other LLMs", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Typical LLM (ChatGPT, Claude, etc.)"
    hdr[1].text = "Spine AI"
    rows = [
        ("Cloud-hosted", "100% local (Ollama on your PC)"),
        ("One model for all tasks", "Task-aware routing — deep model for research/chat, fast model for PC & voice"),
        ("Text in a browser", "Voice + visual orb + system tray"),
        ("No real PC control", "Launches apps, types, organises files, plays music"),
        ("Generic memory", "Your knowledge base (RAG, auto-index)"),
        ("Their privacy policy", "Your data stays on your machine"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_paragraph()

    add_heading(doc, "How it differs from enterprise AI", 2)
    doc.add_paragraph(
        "Enterprise tools are powerful — but they are vendor-locked, cloud-dependent, and built for their ecosystem, "
        "not yours. Spine is the opposite:"
    )
    add_bullets(
        doc,
        [
            "Open stack — Python, Ollama, Whisper, ChromaDB",
            "Runs on consumer hardware (tested on Ryzen 7 + RTX 3050)",
            "You control routing — swap models, bench speed, train a custom brain from your notes",
            "Agentic by design — orchestrator delegates; specialists execute",
        ],
    )
    doc.add_paragraph(
        "This is closer to where agentic AI and edge AI are heading — but shipped as something "
        "one person can run on a laptop today."
    )

    add_heading(doc, "How people can access it", 2)
    doc.add_paragraph("Spine is open source on GitHub:")
    link = doc.add_paragraph()
    run = link.add_run("https://github.com/Sehn1302/Spine_AI")
    run.bold = True
    doc.add_paragraph("Quick start:")
    add_bullets(
        doc,
        [
            "Install Ollama (ollama.com/download)",
            "Pull models: qwen2.5:7b, phi3:mini, nomic-embed-text",
            "Clone the repo, install requirements, run Launch Spine.bat",
        ],
    )
    doc.add_paragraph(
        "Text mode, voice mode, and visual orb mode included. Stress test: Scripts\\run_stress_test.bat"
    )

    add_heading(doc, "How this could revolutionise personal & professional AI", 2)
    doc.add_paragraph(
        "Today, “AI” for most people means asking a cloud chatbot a question. Spine points to a different model:"
    )
    quote = doc.add_paragraph()
    quote.add_run(
        "Your computer becomes the AI platform — private, always-on, multi-skilled, and action-oriented."
    ).italic = True
    doc.add_paragraph("Imagine executives, researchers, and operators with an assistant that:")
    add_bullets(
        doc,
        [
            "Knows their documents (RAG)",
            "Speaks and listens naturally",
            "Does work on the machine — not just suggests it",
            "Uses the right brain for the right job — fast when speed matters, deep when quality matters",
        ],
    )
    doc.add_paragraph(
        "No subscription required. No data exfiltration. No waiting for a vendor to add a feature."
    )

    add_heading(doc, "Future scope", 2)
    add_bullets(
        doc,
        [
            "Vision layer — screen-aware automation (click what it sees)",
            "Cross-device sync — same brain on laptop, desktop, and eventually mobile",
            "Plugin ecosystem — third-party agents (calendar, CRM, code review)",
            "Team mode — shared knowledge base with role-based routing",
            "Mac / Linux ports — same orchestrator, platform-native control layers",
            "Custom model training — models train spine-custom from personal notes",
        ],
    )

    add_heading(doc, "Benchmarks (local, warm runs)", 2)
    bench = doc.add_table(rows=1, cols=2)
    bench.style = "Table Grid"
    bench.rows[0].cells[0].text = "Path"
    bench.rows[0].cells[1].text = "Latency"
    for path, lat in [
        ("Fast model (phi3:mini)", "~0.8s"),
        ("Deep model (qwen2.5:7b)", "~2.1s"),
        ("Whisper STT (GPU vs CPU)", "~0.5s vs ~2s"),
        ("Multi-LLM stress test", "11/11 passed"),
    ]:
        r = bench.add_row().cells
        r[0].text = path
        r[1].text = lat

    doc.add_paragraph()
    doc.add_paragraph(
        "If you are building in agentic AI, local LLMs, voice interfaces, or intelligent automation — "
        "I would love to connect."
    )
    doc.add_paragraph(
        "What would you want a local orchestrator to handle first: voice, files, or research?"
    )

    add_heading(doc, "Hashtags (paste after the post)", 1)
    tags = doc.add_paragraph(
        "#AI #ArtificialIntelligence #LLM #GenAI #GenerativeAI #MachineLearning #DeepLearning #NLP #RAG "
        "#RetrievalAugmentedGeneration #MultiAgent #MultiAgentSystems #AgenticAI #VoiceAI #LocalAI #EdgeAI "
        "#OnDeviceAI #Python #Ollama #Whisper #CUDA #AIProductEngineering #VibeCoding #AIEngineering "
        "#DataAnalytics #Automation #IntelligentAutomation #Innovation #DigitalTransformation "
        "#ProductManagement #OpenToWork #Hiring #TechJobs #SoftwareEngineering #OpenSource #Privacy "
        "#ResponsibleAI #FutureOfWork #AIAutomation #Orchestration #BuildInPublic"
    )
    if tags.runs:
        tags.runs[0].font.size = Pt(10)

    add_heading(doc, "Twitter / X (≤280 characters)", 1)
    doc.add_paragraph(
        "Built Spine AI: local multi-LLM orchestrator — not cloud ChatGPT. Voice, RAG, PC control, "
        "task routing (qwen + phi3). Vibe coding + AI product engineering. Open source: "
        "https://github.com/Sehn1302/Spine_AI #LocalAI #AgenticAI #LLM #VibeCoding"
    )

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
